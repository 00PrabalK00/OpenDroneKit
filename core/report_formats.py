"""PDF and DOCX report writers.

The existing report engine produces HTML and Markdown. Engineering deliverables are
usually asked for as PDF (fixed, signable) or DOCX (editable by the client), so those
are generated directly from the report payload rather than by converting HTML, which
would drag in a browser engine and lose the structure.

Findings are ordered by severity, because a report whose first page buries the
critical item behind twenty cosmetic ones is a worse deliverable regardless of how
complete it is. Numbering stays stable so a finding can be referred to in
correspondence.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

# Ordered worst first. Anything unrecognised sorts last rather than being dropped.
SEVERITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3}


@dataclass
class ReportTemplate:
    """What an organisation wants in its deliverable."""

    title: str = "Inspection Report"
    organization: str = ""
    client: str = ""
    logo_path: str = ""
    author: str = ""
    # Section keys, in the order they should appear.
    sections: list[str] = field(default_factory=lambda: [
        "summary", "methodology", "findings", "measurements", "recommendations",
    ])
    include_unreviewed: bool = True
    severity_order: bool = True

    def to_dict(self) -> dict[str, Any]:
        return {
            "title": self.title, "organization": self.organization, "client": self.client,
            "logo_path": self.logo_path, "author": self.author, "sections": list(self.sections),
            "include_unreviewed": self.include_unreviewed, "severity_order": self.severity_order,
        }


def _severity_rank(finding: dict[str, Any]) -> tuple[int, str]:
    severity = str(finding.get("severity", "")).lower()
    return SEVERITY_ORDER.get(severity, len(SEVERITY_ORDER)), str(finding.get("category", ""))


def prepare_findings(
    findings: list[dict[str, Any]], template: ReportTemplate
) -> list[dict[str, Any]]:
    """Filter and order findings, numbering them stably.

    Numbering happens after ordering so the numbers match the printed sequence, and a
    finding referred to as "3" in correspondence is the third one in the document.
    """
    selected = [
        dict(finding) for finding in findings
        if template.include_unreviewed or finding.get("review_state") == "accepted"
    ]
    if template.severity_order:
        selected.sort(key=_severity_rank)
    for index, finding in enumerate(selected, start=1):
        finding["number"] = index
    return selected


def _measurement_caveat(findings: list[dict[str, Any]]) -> str:
    """State plainly when totals cover only part of what was found."""
    unmeasured = sum(
        1 for finding in findings
        if not finding.get("area_m2") and not finding.get("length_m")
    )
    if not unmeasured:
        return ""
    return (
        f"{unmeasured} of {len(findings)} findings carry no measured extent, so the "
        "area and length totals below describe only the measured ones."
    )


def _unreviewed_caveat(findings: list[dict[str, Any]]) -> str:
    """State when the document contains claims no human has confirmed."""
    model_unreviewed = sum(
        1 for finding in findings
        if finding.get("source") == "model" and finding.get("review_state") != "accepted"
    )
    if not model_unreviewed:
        return ""
    return (
        f"{model_unreviewed} finding(s) in this report were produced by an automated "
        "model and have not been confirmed by an inspector. They are included for "
        "review and should not be treated as verified."
    )


def write_pdf(path: str | Path, payload: dict[str, Any], template: ReportTemplate | None = None) -> str:
    """Render a PDF deliverable."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    template = template or ReportTemplate()
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    caveat_style = ParagraphStyle(
        "Caveat", parent=styles["BodyText"], textColor=colors.HexColor("#8a4b00"),
        backColor=colors.HexColor("#fff4e5"), borderPadding=6, spaceBefore=6, spaceAfter=10,
    )

    document = SimpleDocTemplate(
        str(target), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm, topMargin=18 * mm, bottomMargin=18 * mm,
        title=template.title, author=template.author or template.organization or "OpenDroneKit",
    )

    findings = prepare_findings(payload.get("findings", []), template)
    story: list[Any] = []

    story.append(Paragraph(template.title, styles["Title"]))
    meta = [
        ["Organization", template.organization or "-"],
        ["Client", template.client or payload.get("client", "-") or "-"],
        ["Project", str(payload.get("project_name", "-"))],
        ["Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")],
    ]
    if payload.get("crs_epsg"):
        # A report carrying measurements must say which coordinate system they are in.
        meta.append(["Coordinate system", f"EPSG:{payload['crs_epsg']}"])
    table = Table(meta, colWidths=[45 * mm, 110 * mm])
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f4f4f4")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    story.extend([Spacer(1, 6 * mm), table, Spacer(1, 8 * mm)])

    for caveat in (_unreviewed_caveat(findings), _measurement_caveat(findings)):
        if caveat:
            story.append(Paragraph(caveat, caveat_style))

    if "summary" in template.sections and payload.get("summary"):
        story.append(Paragraph("Executive summary", styles["Heading2"]))
        story.append(Paragraph(str(payload["summary"]), styles["BodyText"]))
        story.append(Spacer(1, 4 * mm))

    if "methodology" in template.sections and payload.get("methodology"):
        story.append(Paragraph("Methodology", styles["Heading2"]))
        story.append(Paragraph(str(payload["methodology"]), styles["BodyText"]))
        story.append(Spacer(1, 4 * mm))

    if "findings" in template.sections:
        story.append(Paragraph(f"Findings ({len(findings)})", styles["Heading2"]))
        if findings:
            rows = [["#", "Category", "Severity", "Area m2", "Length m", "Source", "Review"]]
            for finding in findings:
                rows.append([
                    str(finding["number"]),
                    str(finding.get("category", "-")),
                    str(finding.get("severity", "-")),
                    f"{finding['area_m2']:.3f}" if finding.get("area_m2") else "not measured",
                    f"{finding['length_m']:.2f}" if finding.get("length_m") else "not measured",
                    str(finding.get("source", "-")),
                    str(finding.get("review_state", "-")),
                ])
            findings_table = Table(rows, repeatRows=1,
                                   colWidths=[10 * mm, 35 * mm, 20 * mm, 24 * mm, 24 * mm, 20 * mm, 24 * mm])
            findings_table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef5")),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(findings_table)
        else:
            story.append(Paragraph("No findings were recorded for this inspection.",
                                   styles["BodyText"]))
        story.append(Spacer(1, 6 * mm))

    if "measurements" in template.sections and payload.get("measurements"):
        story.append(Paragraph("Measurements", styles["Heading2"]))
        measurement_rows = [["Quantity", "Value"]]
        for key, value in payload["measurements"].items():
            measurement_rows.append([str(key).replace("_", " "), str(value)])
        measurement_table = Table(measurement_rows, colWidths=[70 * mm, 85 * mm])
        measurement_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef5")),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
        ]))
        story.append(measurement_table)
        story.append(Spacer(1, 6 * mm))

    if "recommendations" in template.sections and payload.get("recommendations"):
        story.append(Paragraph("Recommendations", styles["Heading2"]))
        for item in payload["recommendations"]:
            story.append(Paragraph(f"• {item}", styles["BodyText"]))

    document.build(story)
    return str(target)


def write_docx(path: str | Path, payload: dict[str, Any], template: ReportTemplate | None = None) -> str:
    """Render an editable DOCX deliverable."""
    from docx import Document
    from docx.shared import Pt

    template = template or ReportTemplate()
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)

    findings = prepare_findings(payload.get("findings", []), template)
    document = Document()
    document.core_properties.title = template.title
    document.core_properties.author = template.author or template.organization or "OpenDroneKit"

    document.add_heading(template.title, level=0)

    meta = document.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    entries = [
        ("Organization", template.organization or "-"),
        ("Client", template.client or str(payload.get("client", "-")) or "-"),
        ("Project", str(payload.get("project_name", "-"))),
        ("Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
    ]
    if payload.get("crs_epsg"):
        entries.append(("Coordinate system", f"EPSG:{payload['crs_epsg']}"))
    for label, value in entries:
        row = meta.add_row().cells
        row[0].text = label
        row[1].text = value

    for caveat in (_unreviewed_caveat(findings), _measurement_caveat(findings)):
        if caveat:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(caveat)
            run.bold = True
            run.font.size = Pt(9)

    if "summary" in template.sections and payload.get("summary"):
        document.add_heading("Executive summary", level=1)
        document.add_paragraph(str(payload["summary"]))

    if "methodology" in template.sections and payload.get("methodology"):
        document.add_heading("Methodology", level=1)
        document.add_paragraph(str(payload["methodology"]))

    if "findings" in template.sections:
        document.add_heading(f"Findings ({len(findings)})", level=1)
        if findings:
            table = document.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            header = table.rows[0].cells
            for index, label in enumerate(
                ["#", "Category", "Severity", "Area m2", "Length m", "Source", "Review"]
            ):
                header[index].text = label
            for finding in findings:
                cells = table.add_row().cells
                cells[0].text = str(finding["number"])
                cells[1].text = str(finding.get("category", "-"))
                cells[2].text = str(finding.get("severity", "-"))
                cells[3].text = (f"{finding['area_m2']:.3f}" if finding.get("area_m2")
                                 else "not measured")
                cells[4].text = (f"{finding['length_m']:.2f}" if finding.get("length_m")
                                 else "not measured")
                cells[5].text = str(finding.get("source", "-"))
                cells[6].text = str(finding.get("review_state", "-"))
        else:
            document.add_paragraph("No findings were recorded for this inspection.")

    if "measurements" in template.sections and payload.get("measurements"):
        document.add_heading("Measurements", level=1)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Quantity"
        table.rows[0].cells[1].text = "Value"
        for key, value in payload["measurements"].items():
            cells = table.add_row().cells
            cells[0].text = str(key).replace("_", " ")
            cells[1].text = str(value)

    if "recommendations" in template.sections and payload.get("recommendations"):
        document.add_heading("Recommendations", level=1)
        for item in payload["recommendations"]:
            document.add_paragraph(str(item), style="List Bullet")

    document.save(str(target))
    return str(target)


WRITERS = {"pdf": (write_pdf, ".pdf"), "docx": (write_docx, ".docx")}
