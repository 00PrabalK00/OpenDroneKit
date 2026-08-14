"""PDF and DOCX deliverables.

A non-zero byte count is not evidence that a PDF is valid, so every generated file is
opened with its own reader and its text extracted. The caveat tests matter most: a
report that silently presents unconfirmed model output as findings, or totals over
partly unmeasured defects as the extent of the damage, is misleading even though
every number in it is accurate.
"""

from __future__ import annotations

import pytest

pytest.importorskip("reportlab")
pytest.importorskip("docx")
pypdf = pytest.importorskip("pypdf")

from core.report_formats import (  # noqa: E402
    ReportTemplate,
    prepare_findings,
    write_docx,
    write_pdf,
)

PAYLOAD = {
    "project_name": "Bridge 7",
    "client": "County Roads",
    "crs_epsg": 32617,
    "summary": "Twelve findings recorded across the north elevation.",
    "methodology": "COLMAP reconstruction with RANSAC Helmert georeferencing.",
    "findings": [
        {"category": "efflorescence", "severity": "low", "source": "human",
         "review_state": "accepted", "area_m2": 0.30},
        {"category": "spalling", "severity": "critical", "source": "model",
         "review_state": "accepted", "model_key": "structural", "confidence": 0.91,
         "area_m2": 0.80, "length_m": 1.2},
        {"category": "crack", "severity": "medium", "source": "model",
         "review_state": "unreviewed", "model_key": "crack_segmentation",
         "confidence": 0.62, "length_m": 2.4},
        {"category": "corrosion", "severity": "high", "source": "human",
         "review_state": "accepted"},  # deliberately unmeasured
    ],
    "measurements": {"total_area_m2": 1.10, "surveyed_area_m2": 1600.0},
    "recommendations": ["Inspect the spall before further loading."],
}


def pdf_text(path) -> str:
    reader = pypdf.PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def docx_text(path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class TestFindingOrder:
    def test_findings_are_ordered_worst_first(self):
        """A report that buries the critical item behind cosmetic ones is a worse
        deliverable regardless of how complete it is."""
        ordered = prepare_findings(PAYLOAD["findings"], ReportTemplate())
        assert [f["severity"] for f in ordered] == ["critical", "high", "medium", "low"]

    def test_numbering_follows_the_printed_order(self):
        """A finding referred to as "3" must be the third one in the document."""
        ordered = prepare_findings(PAYLOAD["findings"], ReportTemplate())
        assert [f["number"] for f in ordered] == [1, 2, 3, 4]
        assert ordered[0]["category"] == "spalling"

    def test_unreviewed_can_be_excluded_by_template(self):
        template = ReportTemplate(include_unreviewed=False)
        ordered = prepare_findings(PAYLOAD["findings"], template)
        assert all(f["review_state"] == "accepted" for f in ordered)
        assert len(ordered) == 3

    def test_an_unknown_severity_sorts_last_rather_than_being_dropped(self):
        findings = PAYLOAD["findings"] + [{"category": "odd", "severity": "unheard-of"}]
        ordered = prepare_findings(findings, ReportTemplate())
        assert len(ordered) == 5, "a finding was silently dropped"
        assert ordered[-1]["category"] == "odd"


class TestPdf:
    def test_the_file_parses_as_a_pdf(self, tmp_path):
        """Byte count proves nothing; the file must open with a real reader."""
        target = write_pdf(tmp_path / "report.pdf", PAYLOAD)
        reader = pypdf.PdfReader(target)
        assert len(reader.pages) >= 1

    def test_the_title_and_project_appear(self, tmp_path):
        text = pdf_text(write_pdf(tmp_path / "r.pdf", PAYLOAD,
                                  ReportTemplate(title="Structural Inspection")))
        assert "Structural Inspection" in text
        assert "Bridge 7" in text

    def test_the_coordinate_system_is_stated(self, tmp_path):
        """A report carrying measurements must say which CRS they are in."""
        text = pdf_text(write_pdf(tmp_path / "r.pdf", PAYLOAD))
        assert "32617" in text

    def test_unreviewed_model_findings_are_flagged(self, tmp_path):
        text = pdf_text(write_pdf(tmp_path / "r.pdf", PAYLOAD))
        assert "not been confirmed" in text
        assert "should not be treated as verified" in text

    def test_unmeasured_findings_are_declared(self, tmp_path):
        text = pdf_text(write_pdf(tmp_path / "r.pdf", PAYLOAD))
        assert "no measured extent" in text
        assert "not measured" in text

    def test_a_report_with_no_findings_says_so(self, tmp_path):
        text = pdf_text(write_pdf(tmp_path / "empty.pdf", {"project_name": "Clean site"}))
        assert "No findings were recorded" in text

    def test_a_fully_reviewed_measured_report_carries_no_caveat(self, tmp_path):
        """The warnings must be conditional, not boilerplate on every document."""
        clean = {
            "project_name": "Tidy",
            "findings": [{"category": "crack", "severity": "low", "source": "human",
                          "review_state": "accepted", "area_m2": 0.1, "length_m": 0.5}],
        }
        text = pdf_text(write_pdf(tmp_path / "clean.pdf", clean))
        assert "not been confirmed" not in text
        assert "no measured extent" not in text


class TestDocx:
    def test_the_file_parses_as_a_docx(self, tmp_path):
        from docx import Document

        target = write_docx(tmp_path / "report.docx", PAYLOAD)
        assert len(Document(target).paragraphs) >= 1

    def test_content_and_caveats_appear(self, tmp_path):
        text = docx_text(write_docx(tmp_path / "r.docx", PAYLOAD))
        assert "Bridge 7" in text
        assert "spalling" in text
        assert "not been confirmed" in text
        assert "32617" in text

    def test_findings_land_in_a_table(self, tmp_path):
        from docx import Document

        document = Document(write_docx(tmp_path / "r.docx", PAYLOAD))
        findings_tables = [t for t in document.tables if t.rows[0].cells[0].text == "#"]
        assert findings_tables, "findings were not rendered as a table"
        # One header row plus four findings.
        assert len(findings_tables[0].rows) == 5

    def test_template_section_selection_is_honoured(self, tmp_path):
        template = ReportTemplate(sections=["findings"])
        text = docx_text(write_docx(tmp_path / "r.docx", PAYLOAD, template))
        assert "Executive summary" not in text
        assert "Findings" in text


class TestTemplates:
    def test_organization_and_client_appear(self, tmp_path):
        template = ReportTemplate(organization="Acme Surveying", client="County Roads")
        text = pdf_text(write_pdf(tmp_path / "r.pdf", PAYLOAD, template))
        assert "Acme Surveying" in text
        assert "County Roads" in text

    def test_severity_ordering_can_be_disabled(self):
        template = ReportTemplate(severity_order=False)
        ordered = prepare_findings(PAYLOAD["findings"], template)
        assert ordered[0]["category"] == "efflorescence"  # original input order

    def test_template_round_trips_to_a_dict(self):
        template = ReportTemplate(title="X", organization="Y", sections=["findings"])
        payload = template.to_dict()
        assert payload["title"] == "X"
        assert payload["sections"] == ["findings"]
